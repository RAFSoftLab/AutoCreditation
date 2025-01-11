'''
Reading of .docx files and converting them to .txt files, with conversion of cyrillic characters to latin characters.
'''

import os
import shutil
import docx
# import textract
from doc2docx import convert

import src.cyrillyc_to_latin as cyrillic_to_latin
import src.util as util

# TODO: some files are .doc, check options for conversion to .docx
def convert_doc_to_txt_docx(root_dir, docx_path, file_name, processed_dir='/tmp/converted_documents/', clear_dir=False, data_types=['text', 'tables']):
    '''
    Converts docx to text. By default converts all pages. If conv_all is set to False, it converts only the pages between first_page and last_page.

    Args:
        root_dir (str):          Root directory of the project, absolute path
        docx_path (str):         Relative path to the docx file from the root directory
        file_name (str):         Name of the converted text file (without extension)
        processed_dir (str):     (Optional) Relative path to the directory where the converted txt files will be saved, from the root directory. Default is '/tmp/converted_documents/'
        clear_dir (str):         (Optional) If True, clears the processed_dir directory before converting. Default is False
        data_types (list):       (Optional) List of data types to be extracted from the document. Types are passed as elements of a list. Options are ['text', 'tables']. Default is ['text', 'tables'], which means text and tables are extracted.

    Returns:
        None
    '''


    # Add '/' to start of paths if it is not present
    docx_path = '/{}'.format(docx_path) if docx_path[0] != '/' else docx_path
    processed_dir = '/{}'.format(processed_dir) if processed_dir[0] != '/' else processed_dir

    # Remove processed_dir and all its contents if it exists if clear_dir is set to True
    if clear_dir == True:
        if os.path.exists('{}{}'.format(root_dir, processed_dir)):
            os.system('rm -rf {}'.format('{}{}'.format(root_dir, processed_dir)))
    # Create processed_dir directory
    if not os.path.exists('{}{}'.format(root_dir, processed_dir)):
        os.makedirs('{}{}'.format(root_dir, processed_dir), exist_ok=True)

    # Name of the file to be converted cannot contain spaces, parentheses, or dashes. Replace all special characters with underscores.
    # tmpFile = '{}{}{}'.format(root_dir, processed_dir, docx_path.split('/')[-1].replace(' ', '_').replace('(', '_').replace(')', '_').replace('-', '_'))
    tmpFile = '{}{}{}'.format(root_dir, processed_dir, util.process_name(docx_path.split('/')[-1], file=True))
    shutil.copyfile('{}{}'.format(root_dir, docx_path), tmpFile)

    doc_to_docx_file = ''
    # Convert .doc to .docx if necessary
    # using doc2docx
    if tmpFile[-4:] == '.doc':
        doc_to_docx_file = '/tmp/temp_docx_file.docx'
        # doc_to_docx_file = doc_to_docx_file.replace('/tmp/', '/tmp/converted_documents/')
        if not os.path.exists('{}{}'.format(root_dir, doc_to_docx_file)):
            os.makedirs('{}{}'.format(root_dir, '/'.join(doc_to_docx_file.split('/')[:-1])), exist_ok=True)
        try:
            convert(tmpFile, '{}{}'.format(root_dir, doc_to_docx_file))
        except:
            print('Conversion failed')
            return None
    # using LibreOffice
    # TODO

    # Open .docx file
    doc = docx.Document('{}{}'.format(root_dir, docx_path if doc_to_docx_file == '' else doc_to_docx_file))


    if data_types == ['text', 'tables']:
        #TODO
        print('Whole file extraction:\n')
    if data_types == ['text']:
        # Extract text form file
        print('Text-only extraction:\n')
        newTxtFile_orig = ''
        for paragraph in doc.paragraphs:
            newTxtFile_orig += (paragraph.text + '\n')
            print(paragraph.text)
    elif data_types == ['tables']:
        # Table extraction
        print('Tables-only extraction:\n')
        newTxtFile_orig_tables = '# # # - - - - - - - - - - # # #\n'
        for indexTable, table in enumerate(doc.tables):
            for row in table.rows:
                rowTxt = '|'
                for cell in row.cells:
                    rowTxt += (cell.text + '|')
                print(rowTxt)
                newTxtFile_orig_tables += '{}\n'.format(rowTxt)
            newTxtFile_orig_tables += '# # # - - - - - - - - - - # # #\n'


    # Convert cyrillic characters to latin characters
    newTxtFile = cyrillic_to_latin.cyrillic_to_latin(newTxtFile_orig)

    # Convert cyrillic characters to latin characters in tables
    newTxtFile_tables = cyrillic_to_latin.cyrillic_to_latin(newTxtFile_orig_tables)

    # doc.save method will not save a valid file if cyrillic characters are present,
        # doc.save('{}{}/{}.txt'.format(root_dir, processed_dir, file_name))
    # so text is extracted and saved with python .write method:
    # original file:
    # with open('{}{}/{}.txt'.format(root_dir, processed_dir, '{}_orig'.format(file_name)), 'w', encoding='utf-8') as f:
    #     f.write(newTxtFile_orig)
    # file with cyrillic characters replaced with latin characters:
    with open('{}{}/{}.txt'.format(root_dir, processed_dir, file_name), 'w', encoding='utf-8') as f:
        f.write(newTxtFile)

    # Save tables as .txt file
    with open('{}{}/{}.txt'.format(root_dir, processed_dir, '{}_tables'.format(file_name)), 'w', encoding='utf-8') as f:
        f.write(newTxtFile_tables)

    # Remove temporary file
    os.remove(tmpFile)


# TODO: Implement textract
# def convert_doc_to_txt_textract(root_dir, docx_path, file_name, processed_dir='/tmp/converted_documents/', clear_dir=False):
#     '''

#     '''

#     text = textract.process('{}{}'.format(root_dir, docx_path))
#     print(text)
